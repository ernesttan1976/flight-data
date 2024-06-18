import sys, os
import json 
import datetime
from llama_index.node_parser import SimpleNodeParser
from llama_index.schema import MetadataMode
from llama_index import (
    SimpleDirectoryReader,
    ServiceContext,
)
import streamlit as st
from llama_index.schema import NodeWithScore
import copy 
# from pii2 import NERPIINodePostprocessor
from llama_index.indices.postprocessor import NERPIINodePostprocessor
import random
import re
import pickle
import pandas as pd
import time 
from docx import Document
from python_docx_replace import docx_replace
from dotenv import load_dotenv

load_dotenv()

if 'filename' not in st.session_state:
    st.session_state['filename'] = None

# Set page config to full width and height
st.set_page_config(
    page_title="SwiftScrub",
    page_icon="🫧",
    layout="wide",
)

# Custom CSS to style the scrollable containers
scrollable_css='''
<style>
    section.main>div {
        padding-bottom: 0rem;
    }
    # [data-testid="stVerticalBlock"]>[data-testid="stHorizontalBlock"]:has([data-testid="stMarkdown"]){
    #     overflow: auto;
    #     max-height: 650px;
    # }
    # [data-testid="element-container"] [data-testid="stTable"]{
    #     overflow: auto; 
    #     max-height: 350px;
    # }
    [data-testid="stExpanderDetails"]:has([data-testid="stTable"]){
        overflow: auto; 
        max-height: 350px;
    }
</style>
'''

def check_folder_exists(*folder_names):
    for folder_name in folder_names:
        check_path = os.path.join(os.getcwd(), folder_name)
        if not os.path.exists(check_path):
            os.makedirs(check_path)

def save_original_json(filename, nodes):
    json_original_filename = os.path.join(os.getcwd(), "redacteddocs", f"{os.path.splitext(os.path.basename(filename))[0]}_original.json")
    print(json_original_filename)
    with open(json_original_filename, "w", encoding="utf-8") as file:
        file.write("[\n")
        for i,node in enumerate(nodes):
            node_dict = {
                "text": node.text,
                "metadata": node.metadata,
            }
            print(json.dumps(node_dict,indent=4),file=file)
            if i<len(nodes)-1:
                file.write(",\n")
        file.write("\n]") 

def save_redacted_file(filename, new_nodes):
    # Save new_nodes into a JSON file
    json_filename = os.path.join(os.getcwd(), "redacteddocs", f"{os.path.splitext(os.path.basename(filename))[0]}_redacted.json")
    print(json_filename)
    with open(json_filename, "w", encoding="utf-8") as file:
        file.write("[\n")
        for i,node in enumerate(new_nodes):
            node_dict = {
                "text": node.text,
                "metadata": {
                    **node.metadata,
                    "__pii_node_info__": [],
                }
            }
            print(json.dumps(node_dict,indent=4),file=file)
            if i<len(new_nodes)-1:
                file.write(",\n")
        file.write("\n]") 

def get_metadata():
    meta = json.loads(st.session_state["reverse_merged_dict"])
    return meta

def sort_metadata(reverse_merged_dict):
    sorted_dict = dict(sorted(reverse_merged_dict.items(), key=lambda x: (-len(x[0]), x[0])))
    sorted_dict = {key: value for key, value in sorted_dict.items() if len(key) > 1 and "##" not in key}
    return sorted_dict

def get_merged_dict(new_nodes):
    metadata_array = [node.node.metadata["__pii_node_info__"] for node in new_nodes]
    merged_dict={}
    for metadata_dict in metadata_array:
        merged_dict = dict(merged_dict | metadata_dict)
    print(merged_dict)
    return merged_dict

# if existing metadata exists, to load it, instead of overwriting it. This keeps the memory of the past editing
def save_metadata_file(filename,merged_dict,overwrite=False):
    metadata_filename = os.path.join(os.getcwd(), "redacteddocs", f"{os.path.splitext(os.path.basename(filename))[0]}_metadata.json")
    print(metadata_filename)
    if os.path.exists(metadata_filename) and not overwrite:
        print("Metadata file already exists. Loading metadata into session variable.")
        with open(metadata_filename, "r", encoding="utf-8") as file: 
            reverse_merged_dict = json.load(file)
            st.session_state["reverse_merged_dict"] = json.dumps(sort_metadata(reverse_merged_dict))
        return
    with open(metadata_filename, "w", encoding="utf-8") as file: 
        print(json.dumps(merged_dict, indent=4), file=file)
        st.session_state["reverse_merged_dict"] = json.dumps(sort_metadata(merged_dict))

def save_original_pickle(filename, nodes):
    pickle_original_filename = os.path.join(os.getcwd(), "originaldocs", f"{os.path.splitext(os.path.basename(filename))[0]}_original.pkl")
    print(pickle_original_filename)

    with open(pickle_original_filename, 'wb') as f: # write binary mode
        pickle.dump(nodes, f, protocol=pickle.HIGHEST_PROTOCOL)
        f.close()

def load_original_pickle(filename):
    pickle_original_filename = os.path.join(os.getcwd(), "originaldocs", f"{os.path.splitext(os.path.basename(filename))[0]}_original.pkl")
    print(pickle_original_filename)
    with open(pickle_original_filename, 'rb') as f: # read binary mode
        nodes = pickle.load(f)
        f.close()
    return nodes



def postprocess_nodes_streaming(nodes, processor):
    for node in nodes:
        new_nodes = processor.postprocess_nodes([NodeWithScore(node=node)])
        new_node = new_nodes[0]

        yield new_node

def get_new_nodes_from_metadata(nodes, reverse_merged_dict):
    new_nodes=[]
    for node in nodes:
        new_node = copy.deepcopy(node)
        new_node.text = redact(node.text, reverse_merged_dict)
        new_nodes.append(new_node)
    return new_nodes

def get_metadata_from_file():
    metadata_filename = os.path.join(os.getcwd(), "redacteddocs", f"{os.path.splitext(os.path.basename(st.session_state['filename']))[0]}_metadata.json")
    if os.path.exists(metadata_filename):
        with open(metadata_filename, "r", encoding="utf-8") as file: 
            reverse_merged_dict = json.load(file)
            print("reverse_merged_dict",reverse_merged_dict)
            st.session_state["reverse_merged_dict"] = json.dumps(sort_metadata(reverse_merged_dict))
            return reverse_merged_dict

def process_file(uploaded_file):

    check_folder_exists("originaldocs","redacteddocs","words")
    temp = os.path.join("originaldocs",f"{uploaded_file.name}")
    with open(temp, "wb")  as f:
        f.write(uploaded_file.getbuffer())
        filename = f.name
    if filename == '':
        filename = os.path.join(os.getcwd(),'originaldocs','health_report.pdf')
    console.update(label=f"Loading {filename}", state="running")

    pickle_original_filename = os.path.join(os.getcwd(), "originaldocs", f"{os.path.splitext(os.path.basename(filename))[0]}_original.pkl")
    if os.path.exists(pickle_original_filename):
        nodes = load_original_pickle(filename)
        print(f'Loaded from file and parsed {len(nodes)} nodes')
        console.update(label=f"Loading nodes from file {filename} 🫧", state="running")

        metadata_filename = os.path.join(os.getcwd(), "redacteddocs", f"{os.path.splitext(os.path.basename(filename))[0]}_metadata.json")
        if os.path.exists(metadata_filename):
            with open(metadata_filename, "r", encoding="utf-8") as file: 
                reverse_merged_dict = json.load(file)
                st.session_state["reverse_merged_dict"] = json.dumps(sort_metadata(reverse_merged_dict))

        redactcommonwords(nodes)
        redactmisc(nodes)

        new_nodes = get_new_nodes_from_metadata(nodes, reverse_merged_dict)    
        console.update(label=f"Loaded nodes from file {filename} 😎", state="complete")

    else:
        reader = SimpleDirectoryReader(input_files=[filename])
        docs = reader.load_data()
        print(f'Loaded {len(docs)} docs')

        console.update(label=f"Node parser running for {filename} 🫧", state="running")

        parser = SimpleNodeParser.from_defaults()
        nodes = parser.get_nodes_from_documents(docs, show_progress=True)
        print(f'Parsed {len(nodes)} nodes')

        if len(nodes)==0:
            console.update(label=f"No nodes to process for {filename}", state="error")
            return

        service_context = ServiceContext.from_defaults()
        processor = NERPIINodePostprocessor(service_context=service_context)
        new_nodes_stream = postprocess_nodes_streaming(nodes, processor)

        console.update(label=f"NERPIINodePostprocessor running for {filename} 🫧", state="running")

        nodes = removetabs(nodes)
        new_nodes = []
        total_nodes = 0
        start_time = time.time()

        for new_node in new_nodes_stream:
            total_nodes += 1
            elapsed_time = time.time() - start_time
            remaining_time = elapsed_time / total_nodes * (len(nodes)-total_nodes)
            if total_nodes == len(nodes):
                console.update(label=f"NERPIINodePostprocessor processed {total_nodes} of {len(nodes)} nodes for {filename} 🫧", state="complete")
            else:
                console.update(label=f"NERPIINodePostprocessor processing {total_nodes} of {len(nodes)} nodes for {filename} 🫧", state="running")
            st.session_state["progress"] = float(total_nodes) / float(len(nodes))
            progress_bar.progress(value=float(st.session_state["progress"]), text=f"Time elapsed: {elapsed_time:.0f}s, Remaining time: {remaining_time:.0f}s")
            new_nodes.append(new_node)


        node_count = len(nodes)
        print(len(new_nodes))
        # For Logging and Data  
        meta={}
        for page in range(node_count):
            meta = meta | new_nodes[page].node.metadata["__pii_node_info__"]

        # Save nodes into a JSON file - for logging purpose
        save_original_json(filename, nodes)

        # Save nodes into a PKL pickle file, for rapid object load/unload
        save_original_pickle(filename, nodes)

        # Save new_nodes into a JSON file
        save_redacted_file(filename, new_nodes)
        
        # get_merged_dict
        reverse_merged_dict = reverse_dictionary(get_merged_dict(new_nodes))


        # if existing metadata exists, to load it, instead of overwriting it. This keeps the memory of the past editing
        # Save node.node.metadata["pii_node_info"] into a JSON file
        save_metadata_file(filename, reverse_merged_dict, overwrite=False)

        redactcommonwords(nodes)

        # use regex to redact the emails, urls, phone_numbers document_numbers, addresses
        redactmisc(nodes)
    return len(nodes)

def handlereapplyclick():
    st.session_state['file_rerun']=True

def reverse_dictionary(original_dict):
    reversed_dict = {}
    for key, value in original_dict.items():
        reversed_dict[value] = key
    sorted_dict = dict(sorted(reversed_dict.items(), key=lambda x: (-len(x[0]), x[0])))
    sorted_dict = {key: value for key, value in sorted_dict.items() if len(key) > 1 and "##" not in key}

    return sorted_dict

def replace_ignorecase(text, key, value):
    compiled = re.compile(re.escape(key), flags=re.IGNORECASE)
    result = compiled.sub(value, text)
    return result

def redact(old_text, meta_json):
    if type(meta_json) == "str":
        print("error in redact, meta_json is a string")
        print(meta_json)
        return
    text = old_text
    # This updated function will skip replacing any keys that start with "[MISC_", "[ORG_", "[PER_", or "[LOC_".
    for key, value in meta_json.items():
        if key.startswith("[MISC_") or key.startswith("[ORG_") or key.startswith("[PER_") or key.startswith("[LOC_"):
            continue
        text = replace_ignorecase(text, key, value)
    return text

def update_new_nodes(filename):
    nodes = load_original_pickle(filename)
    meta_json = get_metadata()
    new_nodes = []
    for node in nodes:
        new_node = copy.deepcopy(node)
        new_node.text = redact(node.text, meta_json)    
        new_nodes.append(new_node)

    save_redacted_file(filename, new_nodes)
    save_metadata_file(filename, meta_json, overwrite=True)
    return

def handlefilechange():
    st.session_state["reverse_merged_dict"]=None
    st.session_state["file_run_clicked"]=False
    st.session_state["file_run_done"]=False
    st.session_state["file_rerun"]=False
    st.session_state["file_rerun_done"]=False
    st.session_state["file_export_done"]=False
    st.session_state['file_loaded'] = True

           
def handlemetadatachange():
    st.session_state['file_rerun'] = False
    st.session_state['file_rerun_done']=False

def handlerunclick():
    st.session_state['file_run_clicked'] = True

def handlererun():
    st.session_state['file_rerun'] = True
    st.session_state['file_rerun_done'] = False

def removetabs(nodes):
    pattern = re.compile('\\n\\t(?!\\t)') # more efficient
    new_nodes=[]
    for node in nodes:
        new_node = copy.deepcopy(node)
        new_node.text = pattern.sub('',node.text)
        new_nodes.append(new_node)
    return new_nodes

def redactmisc(nodes):
    reverse_merged_dict = get_metadata()

    redact_filename = os.path.join(os.getcwd(), "words", "redact_regex.json")
    print(redact_filename)
    if os.path.exists(redact_filename):
        with open(redact_filename, "r", encoding="utf-8") as file: 
            rd = json.load(file)

            for key in rd:
                for node in nodes:
                    found = re.findall(key['pattern'], node.text)
                    print(key["key"],found)
                    for index, item in enumerate(found):
                        if ":" in item:
                            found.pop(index)
                            print(item, " is invalid because it contains a colon")

                    key["data"].extend(found)
                
            for key in rd:
                for item in key['data']:
                    reverse_merged_dict = {**reverse_merged_dict} | {item: f"[{key['tag']}_{random.randrange(1000, 9999, 1)}]"}

            sorted_metadata = sort_metadata(reverse_merged_dict)
            st.session_state["reverse_merged_dict"] = json.dumps(sort_metadata(reverse_merged_dict))
            print([f"{item['key']} = {item['data']}" for item in rd])

            save_metadata_file(st.session_state["filename"],sorted_metadata,overwrite=True)


def findtextinnodes(text, nodes):
    for node in nodes:
        if text in node.text:
            return True
    return False

def redactcommonwords(nodes):
    metadata = get_metadata()
    add_filename = os.path.join(os.getcwd(), "words", "add_words.json")
    print(add_filename)
    if os.path.exists(add_filename):
        with open(add_filename, "r", encoding="utf-8") as file: 
            add_dict = json.load(file)
            for key in list(add_dict.keys()):
                if findtextinnodes(key, nodes):
                    metadata = {
                        **metadata, key: add_dict[key]
                    }

    delete_filename = os.path.join(os.getcwd(), "words", "delete_words.json")
    print(delete_filename)
    if os.path.exists(delete_filename):
        with open(delete_filename, "r", encoding="utf-8") as file: 
            delete_dict = json.load(file)
            for key in delete_dict:
                if key in metadata:
                    del metadata[key]
    st.session_state["reverse_merged_dict"] = json.dumps(sort_metadata(metadata))

def wordslist(text, tag, type="add"):
    if type=="add" or type =="delete":
        metadata_filename = os.path.join(os.getcwd(), "words", f"{type}_words.json")
        print(metadata_filename)
        if os.path.exists(metadata_filename):
            with open(metadata_filename, "r", encoding="utf-8") as file: 
                metadata = json.load(file)
            with open(metadata_filename, "w", encoding="utf-8") as file:
                new_metadata = {
                    **metadata, text: tag
                }
                sorted_dict = dict(sorted(new_metadata.items(), key=lambda x: (-len(x[0]), x[0])))
                sorted_dict = {key: value for key, value in sorted_dict.items() if len(key) > 1 and "##" not in key}
                print(json.dumps(sorted_dict, indent=4), file=file)
        else:
            with open(metadata_filename, "w", encoding="utf-8") as file:
                new_metadata = {
                    text: tag
                }
                print(json.dumps(new_metadata, indent=4), file=file)


def get_sorted_metadata(reverse_merged_dict):
    if reverse_merged_dict == None:
        return
    sorted_list = [(key, value) for key, value in reverse_merged_dict.items()]
    sorted_list.sort(key=lambda elem: (-len(elem[0]), elem[0]))
    return sorted_list


def addedittag(text2, tag2):
    text = text2.strip()
    tag = f"[{tag2}_{random.randrange(1000, 9999, 1)}]"
    reverse_merged_dict = get_metadata()
    reverse_merged_dict = {**reverse_merged_dict, text: tag}

    st.session_state["reverse_merged_dict"] = json.dumps(sort_metadata(reverse_merged_dict))
    st.session_state['file_rerun'] = True
    st.session_state['file_rerun_done'] = False
    redact_text=""
    wordslist(text, tag, type="add")
    save_metadata_file(st.session_state['filename'],reverse_merged_dict,overwrite=True)

    
def deletetag(redact_text2):
    redact_text = redact_text2.strip()
    reverse_merged_dict = get_metadata()

    reverse_merged_dict.pop(redact_text)
    st.session_state["reverse_merged_dict"] = json.dumps(sort_metadata(reverse_merged_dict))
    st.session_state['file_rerun'] = True
    st.session_state['file_rerun_done'] = False
    wordslist(redact_text, "" , type="delete")
    save_metadata_file(st.session_state['filename'],reverse_merged_dict,overwrite=True)

def get_new_nodes(nodes, reverse_merged_dict):
    new_nodes=[]
    for node in nodes:
        new_node = copy.deepcopy(node)
        new_node.text = redact(node.text, reverse_merged_dict)
        new_nodes.append(new_node)
    return new_nodes


def handlegenerate():
    st.session_state['file_export']=True
    st.session_state['file_export_done']=False

def export_redacted_file():
    sorted_metadata = get_sorted_metadata(json.loads(st.session_state['reverse_merged_dict']))
    if sorted_metadata == None:
        print("Error, no metadata")
        st.session_state['file_export']=False
        st.session_state['file_export_done']=False
        return False

    # if os.path.splitext(os.path.basename(st.session_state['filename']))[1]==".pdf":
        # disables until alternative to aspose pdf is found

        # input_filename = os.path.join(os.getcwd(), "originaldocs", f"{os.path.splitext(os.path.basename(st.session_state['filename']))[0]}.pdf")
        # output_filename = os.path.join(os.getcwd(), "redacteddocs", f"{os.path.splitext(os.path.basename(st.session_state['filename']))[0]}_redacted.pdf")   
        # # document = ap.Document(input_filename)
        # for (key, value) in sorted_metadata:
        #     print(key,value)
        #     # txtAbsorber = ap.text.TextFragmentAbsorber(key)
        #     document.pages.accept(txtAbsorber)
        #     textFragmentCollection = txtAbsorber.text_fragments
        #     for txtFragment in textFragmentCollection:
        #         txtFragment.text = value        
        # document.save(output_filename)
        # st.session_state['file_export']=False
        # st.session_state['file_export_done']=True
        # return True
    if os.path.splitext(os.path.basename(st.session_state['filename']))[1]==".docx":
        input_filename = os.path.join(os.getcwd(), "originaldocs", f"{os.path.splitext(os.path.basename(st.session_state['filename']))[0]}.docx")
        output_filename = os.path.join(os.getcwd(), "redacteddocs", f"{os.path.splitext(os.path.basename(st.session_state['filename']))[0]}_redacted.docx")   
        doc = Document(input_filename)

        for (key, value) in sorted_metadata:
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text:
                        replaced_text = re.sub(key,value,run.text, 999)
                        if replaced_text!= run.text:
                            run.text=replaced_text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text:
                                    replaced_text = re.sub(key,value,run.text, 999)
                                    if replaced_text!= run.text:
                                        run.text=replaced_text

        doc.save(output_filename)
        st.session_state['file_export']=False
        st.session_state['file_export_done']=True
        return True
    

if 'file_loaded' not in st.session_state:
    st.session_state['file_loaded'] = False

if 'file_run_clicked' not in st.session_state:
    st.session_state['file_run_clicked'] = False

if 'file_run_done' not in st.session_state:
    st.session_state['file_run_done'] = False

if 'enable_reapply_button' not in st.session_state:
    st.session_state['enable_reapply_button'] = False

if 'file_rerun' not in st.session_state:
    st.session_state['file_rerun'] = False

if 'file_rerun_done' not in st.session_state:
    st.session_state['file_rerun_done'] = False

if 'file_export' not in st.session_state:
    st.session_state['file_export'] = False

if 'file_export_done' not in st.session_state:
    st.session_state['file_export_done'] = False

if 'progress' not in st.session_state:
    st.session_state['progress'] = 0 

c1left, c1right= st.container().columns([0.7,0.3])
c1left.markdown("### SwiftScrub🦙&nbsp;🧼&nbsp;:sunglasses:&nbsp;🫧&nbsp;🪪")
h1 = c1right.container().expander(label="Load File", expanded=not st.session_state["file_run_done"])
uploaded_file = h1.file_uploader("File Dialog", on_change=handlefilechange, label_visibility="collapsed")
progress_bar = h1.progress(value=0, text="")
console_container = c1right.container()
console = console_container.status(label="Please load file 📃", state="error")

c3 = c1left.container()
col3 = c1right

if not uploaded_file==None and not st.session_state['file_run_done']:
    st.session_state.filename = uploaded_file.name
    console.update(label=f"Ready... 🦙🫧", state="complete")
if uploaded_file==None:
    st.session_state["file_loaded"]=False

if st.session_state["file_loaded"]: 
    h1.button("Run 🧼", disabled=not st.session_state["file_loaded"], on_click=handlerunclick, use_container_width=True)

if st.session_state['file_run_clicked']: 
    if not st.session_state['file_run_done']:
        st.session_state['node_count'] = process_file(uploaded_file)
        st.session_state['file_run_done'] = True


    if st.session_state['node_count']==0:
        console.update(label='No nodes were readable', state="error") 
    
    if st.session_state['file_run_done'] or st.session_state['file_rerun_done'] and st.session_state['node_count']>0:
        nodes = load_original_pickle(st.session_state['filename'])
        reverse_merged_dict = get_metadata()
        tablabels = [f"pg{page+1}" for page in range(len(nodes))]
        c3left, c3right = c3.columns(2)
        c3left.markdown("##### Original")
        c3right.markdown("##### Redacted")
        tabs = c3.tabs(tablabels)
        new_nodes = get_new_nodes(nodes, reverse_merged_dict)

        for page in range(len(nodes)):
            col1_column,col2_column = tabs[page].columns(2)
            col1_container = col1_column.container()
            col1 = col1_container.container()
            # markdown
            col1.write(nodes[page].text)
            col1.markdown(scrollable_css, unsafe_allow_html=True)
            col2_container = col2_column.container()
            col2 = col2_container.container()
            # markdown
            col2.write(new_nodes[page].text)
            col2.markdown(scrollable_css, unsafe_allow_html=True)

        if st.session_state['file_rerun_done']:
            st.session_state['file_rerun']=False


if st.session_state['file_loaded'] and st.session_state['file_run_done']:

    metatextareacontainer = col3.expander(label="Metadata", expanded=True).container()


    reverse_merged_dict = get_metadata()
    metadata_key_list= list(reverse_merged_dict.keys())
    metadata_list = [[key, reverse_merged_dict[key]] for key in metadata_key_list]      
    df = pd.DataFrame(metadata_list, columns=("Text","Tag"))
    metatextareacontainer.table(df)
    
    

    col3_expander = col3.expander(label="Redact", expanded=True)
    tag1, tag2, tag3 = col3_expander.container().columns(spec=[0.45,0.30,0.25])
    redact_text = tag1.text_input(label="Redact Text", value="")
    redact_tag = tag2.selectbox(label="Tag",options=["ORG","LOC","PER","MISC"])
    tag3.write(" ")
    tag3.write(" ")
    tag3.button(label="Add Tag", on_click=addedittag, args=(redact_text, redact_tag))
    tag4, tag5 = col3_expander.container().columns(spec=[0.7,0.3])
    delete_redact_text = tag4.selectbox(label="Redact Text",options=reverse_merged_dict.keys(), index=0)
    tag5.write(" ")
    tag5.write(" ")
    tag5.button(label="Delete Tag",on_click=deletetag, args=(delete_redact_text,))

    buttonrow = col3_expander.container()
    button_0, button_1, button_2, button_3 = buttonrow.columns(4)
    
    button_0.button(key="button0", label="Generate", on_click=handlegenerate)

    redacted_json = os.path.join(os.getcwd(), "redacteddocs", f"{os.path.splitext(os.path.basename(st.session_state['filename']))[0]}_redacted.json")
    print("redacted_json",redacted_json)        
    with open(redacted_json, "r", encoding="utf-8") as file:            
        json_data = json.dumps(json.load(file), indent=4)
        button_2.download_button(key="button2",label="Redacted .json", data=json_data, file_name=f"{os.path.splitext(os.path.basename(st.session_state['filename']))[0]}_redacted.json")

    if st.session_state['file_export']==True and st.session_state['file_export_done']==False:
        export_redacted_file()
        st.session_state['file_export_done']=True
        st.session_state['file_export']==False


    redacted_file = os.path.join(os.getcwd(), "redacteddocs", f"{os.path.splitext(os.path.basename(st.session_state['filename']))[0]}_redacted{os.path.splitext(os.path.basename(st.session_state['filename']))[1]}")
    if os.path.exists(redacted_file) and st.session_state['file_export_done']:
        with open(redacted_file, "rb") as input_file:
            content = input_file.read()
            button_1.download_button(key="button1",label=f"Redacted {os.path.splitext(os.path.basename(st.session_state['filename']))[1]}", data=content, file_name=f"{os.path.splitext(os.path.basename(st.session_state['filename']))[0]}_redacted{os.path.splitext(os.path.basename(st.session_state['filename']))[1]}", use_container_width=True)
        
    metadata_filename = os.path.join(os.getcwd(), "redacteddocs", f"{os.path.splitext(os.path.basename(st.session_state['filename']))[0]}_metadata.json")
    print("metadata_filename",metadata_filename)
    with open(metadata_filename, "r", encoding="utf-8") as file:
        json_data = json.dumps(json.load(file), indent=4)
        button_3.download_button(key="button3",label="Metadata .json", data=json_data, file_name=f"{os.path.splitext(os.path.basename(st.session_state['filename']))[0]}_metadata.json", mime='application/json', use_container_width=True)

    if st.session_state['file_rerun'] and not st.session_state['file_rerun_done']:
        console.update(label=f"Running redaction again... 🦙🫧", state="running")
        update_new_nodes(st.session_state["filename"])
        st.session_state['file_rerun_done']=True
    
    if st.session_state['file_rerun_done']:
        console.update(label='Redaction complete 🫧🫧🫧', state="complete")


